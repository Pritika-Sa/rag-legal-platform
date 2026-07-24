import os
import re
import json
import hashlib
from datetime import datetime
from typing import List, Dict, Any, Optional
import pdfplumber
import pytesseract
from PIL import Image
from docx import Document
from pydantic import BaseModel, Field
from langchain_core.tools import BaseTool
from langchain_text_splitters import RecursiveCharacterTextSplitter

IMAGE_EXTENSIONS = (".png", ".jpg", ".jpeg")

if os.getenv("TESSERACT_CMD"):
    pytesseract.pytesseract.tesseract_cmd = os.getenv("TESSERACT_CMD")


def extract_text_from_image(file_path: str) -> str:
    """OCRs an image (PNG/JPG/JPEG) into raw text via Tesseract. Raises a
    plain RuntimeError with a user-facing message if the Tesseract binary
    itself isn't installed/on PATH, so callers can show a clean warning
    instead of a raw pytesseract traceback."""
    try:
        return pytesseract.image_to_string(Image.open(file_path))
    except pytesseract.TesseractNotFoundError:
        raise RuntimeError(
            "OCR engine not found. Install Tesseract-OCR and, if it isn't on your system PATH, "
            "set the TESSERACT_CMD environment variable to its executable path."
        )

# Input Schema for LangChain Component
class DocumentParserInput(BaseModel):
    file_path: str = Field(description="The absolute path to the PDF or DOCX file to be parsed")
    document_id: Optional[str] = Field(None, description="Optional unique identifier for the document. If not provided, an MD5 hash of the file will be generated.")
    version: Optional[int] = Field(1, description="Document version number")

# Reusable LangChain Tool Component
class DocumentParsingTool(BaseTool):
    name: str = "document_parser"
    description: str = "Parses PDF and DOCX files, extracts metadata, chunks text using RecursiveCharacterTextSplitter, and returns structured JSON."
    args_schema: type[BaseModel] = DocumentParserInput

    def _run(self, file_path: str, document_id: Optional[str] = None, version: int = 1) -> str:
        """Runs document parsing, chunking, and returns structured JSON output."""
        try:
            result = parse_document_to_json(file_path, document_id, version)
            return json.dumps(result, indent=2)
        except Exception as e:
            return json.dumps({"error": str(e)})

def get_file_md5(file_path: str) -> str:
    """Computes MD5 hash of a file."""
    hasher = hashlib.md5()
    with open(file_path, 'rb') as f:
        buf = f.read()
        hasher.update(buf)
    return hasher.hexdigest()

def extract_pdf_metadata(file_path: str) -> Dict[str, Any]:
    """Extracts metadata and statistics from PDF file."""
    meta = {}
    try:
        with pdfplumber.open(file_path) as pdf:
            meta["page_count"] = len(pdf.pages)
            # Fetch default metadata properties if present
            if pdf.metadata:
                for k, v in pdf.metadata.items():
                    if isinstance(v, (str, int, float, bool)):
                        meta[k.lower()] = v
    except Exception as e:
        meta["error"] = f"Failed to extract PDF metadata: {str(e)}"
    return meta

def extract_docx_metadata(file_path: str) -> Dict[str, Any]:
    """Extracts metadata and properties from Word document."""
    meta = {}
    try:
        doc = Document(file_path)
        props = doc.core_properties
        meta["author"] = props.author or "Unknown"
        meta["title"] = props.title or "Unknown"
        meta["created"] = str(props.created) if props.created else None
        meta["modified"] = str(props.modified) if props.modified else None
        meta["revision"] = props.revision
    except Exception as e:
        meta["error"] = f"Failed to extract DOCX metadata: {str(e)}"
    return meta

def parse_document_to_json(file_path: str, document_id: Optional[str] = None, version: int = 1) -> Dict[str, Any]:
    """Extracts text and constructs structured JSON chunks and metadata."""
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
        
    doc_name = os.path.basename(file_path)
    ext = os.path.splitext(file_path)[1].lower()
    
    # 1. Compute/Assign Document ID and basic info
    if not document_id:
        document_id = get_file_md5(file_path)
        
    upload_date = datetime.now().isoformat()
    
    # 2. Extract text and metadata
    raw_text = ""
    doc_metadata = {
        "file_size_bytes": os.path.getsize(file_path),
        "file_type": ext.replace(".", "")
    }
    
    if ext == ".docx":
        doc = Document(file_path)
        raw_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        doc_metadata.update(extract_docx_metadata(file_path))
    elif ext == ".pdf":
        with pdfplumber.open(file_path) as pdf:
            text_blocks = []
            for page in pdf.pages:
                text_blocks.append(page.extract_text() or "")
            raw_text = "\n".join(text_blocks)
        doc_metadata.update(extract_pdf_metadata(file_path))
    elif ext == ".txt":
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            raw_text = f.read()
    elif ext in IMAGE_EXTENSIONS:
        raw_text = extract_text_from_image(file_path)
    else:
        raise ValueError(f"Unsupported file format: {ext}")
        
    # 3. Create Chunks using RecursiveCharacterTextSplitter
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=150,
        length_function=len
    )
    
    chunks = text_splitter.split_text(raw_text)
    
    # 4. Format structured chunks
    formatted_chunks = []
    for idx, chunk in enumerate(chunks):
        formatted_chunks.append({
            "chunk_id": f"{document_id}_chunk_{idx}",
            "text_content": chunk,
            "metadata": {
                "chunk_index": idx,
                "length": len(chunk)
            }
        })
        
    # 5. Compile structured JSON output representation
    return {
        "document_id": document_id,
        "document_name": doc_name,
        "upload_date": upload_date,
        "version": version,
        "document_metadata": doc_metadata,
        "chunks": formatted_chunks
    }


# Backwards compatibility layer for local orchestrator / regex parsing

# Tried in order, first match wins. Real headings are short, so the shared
# length guard at each call site shrank from 150 to 80 chars — the looser
# patterns below (ALL-CAPS, Title-Case) need the tighter bound to avoid
# matching body text. Lettered sub-headings like "(a)", "(b)" are
# deliberately NOT treated as section boundaries: they're sub-points within
# one clause, not new clauses, and treating them as boundaries would
# fragment clauses instead of fixing under-segmentation.
SECTION_PATTERNS = [
    # Numbered: "Section 3.2", "Article IV.", "1. Foo" (original pattern)
    re.compile(r'^(?:section|clause|article|part)\s+\d+(?:\.\d+)*[:\-\s\.]|^\d+\.\s+[A-Z]', re.IGNORECASE),
    # Roman numerals: "III. Termination", "Article IV - Payment"
    re.compile(r'^(?:article\s+)?[IVXLCDM]{1,6}\.?\s*[-:.]?\s*[A-Z]', re.IGNORECASE),
    # ALL-CAPS heading on its own line, e.g. "TERMINATION", "GOVERNING LAW"
    # (requires zero lowercase letters, which ordinary body sentences always
    # have, so this doesn't misfire on paragraph text)
    re.compile(r'^[A-Z][A-Z0-9 &,\-]{2,59}$'),
    # Un-numbered legal title in Title Case, e.g. "Governing Law",
    # "Confidentiality Obligations" (every word must start uppercase, which
    # ordinary sentences fail on their lowercase function words)
    re.compile(r'^[A-Z][a-zA-Z]*(?:\s+[A-Z][a-zA-Z]*){0,5}$'),
]

SECTION_HEADING_MAX_CHARS = 80

# Title-Case headings that include a lowercase connector word ("Conditions
# of Coverage", "Notice of Cancellation", "Limitation of Use", "Waiver of
# Subrogation") fail SECTION_PATTERNS[3] above because it requires every
# word to start uppercase -- extremely common in insurance-policy section
# titles specifically. Checked word-by-word instead of one regex so the
# connector-word set stays easy to extend.
_TITLE_CASE_CONNECTOR_WORDS = {
    "of", "and", "the", "in", "for", "to", "on", "by", "with", "or",
    "a", "an", "as", "under", "per", "at", "from",
}


def _is_title_case_with_connectors(line: str) -> bool:
    words = line.split()
    if not (2 <= len(words) <= 8):
        return False
    if line.endswith((".", ",", ";")):
        return False
    if not words[0][:1].isupper() or not words[-1][:1].isupper():
        return False
    for word in words:
        cleaned = word.strip(":-").replace("'", "")
        if not cleaned or not cleaned.isalpha():
            return False
        if cleaned.lower() in _TITLE_CASE_CONNECTOR_WORDS:
            continue
        if not cleaned[0].isupper():
            return False
    return True


def _is_section_heading(line: str) -> bool:
    if len(line) >= SECTION_HEADING_MAX_CHARS:
        return False
    return any(p.match(line) for p in SECTION_PATTERNS) or _is_title_case_with_connectors(line)


def _table_row_to_text(header: List[Optional[str]], row: List[Optional[str]]) -> str:
    """Turns one table row into a readable clause-candidate sentence, e.g.
    header ["Coverage", "Sum Insured", "Premium"] + row ["Own Damage",
    "500000", "12000"] -> "Coverage: Own Damage. Sum Insured: 500000.
    Premium: 12000." Falls back to plain cell-joining when there's no usable
    header (single-column schedules, label/value pairs with no header row)."""
    has_header = bool(header) and any(h and str(h).strip() for h in header)
    if has_header:
        pairs = [
            f"{_strip_cid_garbage(str(h)).strip()}: {_strip_cid_garbage(str(c)).strip()}"
            for h, c in zip(header, row)
            if h and str(h).strip() and c and str(c).strip()
        ]
        pairs = [p for p in pairs if p.strip(": .")]
        if pairs:
            return ". ".join(pairs) + "."
    cells = [_strip_cid_garbage(str(c)).strip() for c in row if c and str(c).strip()]
    return "; ".join(c for c in cells if c)


def _extract_page_tables(page, page_num: int) -> tuple:
    """Finds every table on a pdfplumber page and returns
    (table_sections, table_bboxes): one clause-candidate section PER ROW
    (per the requirement that table rows -- not whole tables -- become
    clause text), plus each table's bounding box so the caller can exclude
    that region from the page's plain-text extraction and avoid extracting
    the same content twice. Cells are extracted with a tight x_tolerance for
    the same reason _reconstruct_text_from_words exists below: pdfplumber's
    default cell-text extraction glues adjacent words together on PDFs with
    unusual glyph spacing (observed on real insurer-issued policy PDFs),
    e.g. "RegistrationNumber" instead of "Registration Number"."""
    table_sections: List[Dict[str, Any]] = []
    table_bboxes: List[tuple] = []
    try:
        tables = page.find_tables()
    except Exception:
        return table_sections, table_bboxes

    for t_idx, table in enumerate(tables):
        try:
            data = table.extract(x_tolerance=1)
        except Exception:
            continue
        if not data:
            continue
        table_bboxes.append(table.bbox)

        header = data[0] if len(data) > 1 else []
        has_header = len(data) > 1 and any(h and str(h).strip() for h in header)
        data_rows = data[1:] if has_header else data

        for row in data_rows:
            text = _table_row_to_text(header if has_header else [], row)
            if len(text) < 8:
                continue
            table_sections.append({
                "section_name": f"Policy Schedule Table {t_idx + 1} (Page {page_num})",
                "text_content": text,
                "page_num": page_num,
            })
    return table_sections, table_bboxes


# Unmapped glyph placeholders pdfplumber emits for characters it can't
# resolve to Unicode (fonts with no usable ToUnicode CMap -- seen on real
# insurer-issued PDFs, typically in decorative header glyphs). Pure noise:
# stripped so it doesn't pollute heading/keyword matching on the readable
# text immediately following it on the same line.
_CID_GARBAGE_RE = re.compile(r'(?:\(cid:\d+\))+')


def _strip_cid_garbage(text: str) -> str:
    return _CID_GARBAGE_RE.sub('', text)


def _reconstruct_text_from_words(words: List[dict]) -> str:
    """Rebuilds line-broken text from pdfplumber word boxes instead of
    page.extract_text(), which on some PDFs (unusual glyph spacing, seen on
    real insurer-issued policy PDFs) glues adjacent words together with no
    space at all -- e.g. "RegisteredandHeadOffice:BajajAllianzHouse" --
    silently defeating almost every downstream keyword/regex match.
    extract_words() detects word boundaries from the actual glyph gaps,
    independent of that heuristic, so grouping those boxes back into lines
    by vertical position recovers proper spacing."""
    if not words:
        return ""
    ordered = sorted(words, key=lambda w: (round(w["top"]), w["x0"]))
    lines: List[List[dict]] = []
    current_top = None
    current_words: List[dict] = []
    for w in ordered:
        top = w["top"]
        if current_top is not None and abs(top - current_top) > 3:
            lines.append(current_words)
            current_words = []
            current_top = None
        if current_top is None:
            current_top = top
        current_words.append(w)
    if current_words:
        lines.append(current_words)

    text_lines = []
    for line_words in lines:
        line_words.sort(key=lambda w: w["x0"])
        line_text = _strip_cid_garbage(" ".join(w["text"] for w in line_words)).strip()
        if line_text:
            text_lines.append(line_text)
    return "\n".join(text_lines)


def _page_text_excluding_tables(page, table_bboxes: List[tuple]) -> str:
    """Extracts a page's readable text with table regions filtered out (so
    table content isn't captured twice -- once here, once cleanly via
    _extract_page_tables) and word-glyph gaps reconstructed into real
    spaces via _reconstruct_text_from_words. Falls back to plain
    page.extract_text() if word-level extraction raises or yields nothing,
    so this can only improve extraction quality, never regress it."""
    target_page = page
    if table_bboxes:
        def _outside_all_tables(obj):
            return not any(
                bbox[0] <= obj["x0"] and obj["x1"] <= bbox[2] and bbox[1] <= obj["top"] and obj["bottom"] <= bbox[3]
                for bbox in table_bboxes
            )
        try:
            target_page = page.filter(_outside_all_tables)
        except Exception:
            target_page = page

    try:
        words = target_page.extract_words(x_tolerance=1, y_tolerance=3)
        text = _reconstruct_text_from_words(words)
        if text:
            return text
    except Exception:
        pass

    try:
        return target_page.extract_text() or ""
    except Exception:
        return page.extract_text() or ""


def parse_document(file_path: str) -> List[Dict[str, Any]]:
    """Legacy backward-compatible parser segmenting text into sections using regex or fallback chunks."""
    ext = os.path.splitext(file_path)[1].lower()

    # Fallback to structured chunking if regex segmentation yields fewer than 3 segments
    try:
        sections = []
        if ext == ".docx":
            doc = Document(file_path)
            current_section = "Preamble"
            current_content = []
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if not text:
                    continue
                if _is_section_heading(text):
                    if current_content:
                        sections.append({
                            "section_name": current_section,
                            "text_content": "\n".join(current_content).strip(),
                            "page_num": None
                        })
                    current_section = text
                    current_content = []
                else:
                    current_content.append(text)
            if current_content:
                sections.append({
                    "section_name": current_section,
                    "text_content": "\n".join(current_content).strip(),
                    "page_num": None
                })
        elif ext == ".pdf":
            current_section = "Preamble"
            current_content = []
            current_page = 1
            with pdfplumber.open(file_path) as pdf:
                for page_idx, page in enumerate(pdf.pages):
                    page_num = page_idx + 1
                    # Tables (policy schedules, coverage/premium breakdowns, etc.)
                    # are pulled out row-by-row *before* plain-text extraction,
                    # and their regions excluded from that extraction, so table
                    # content is captured cleanly exactly once instead of being
                    # flattened/garbled into ordinary prose lines or skipped.
                    table_sections, table_bboxes = _extract_page_tables(page, page_num)
                    text = _page_text_excluding_tables(page, table_bboxes)
                    if text:
                        for line in text.split("\n"):
                            line_str = line.strip()
                            if not line_str:
                                continue
                            if _is_section_heading(line_str):
                                if current_content:
                                    sections.append({
                                        "section_name": current_section,
                                        "text_content": "\n".join(current_content).strip(),
                                        "page_num": current_page
                                    })
                                current_section = line_str
                                current_content = []
                                current_page = page_num
                            else:
                                current_content.append(line_str)
                    sections.extend(table_sections)
            if current_content:
                sections.append({
                    "section_name": current_section,
                    "text_content": "\n".join(current_content).strip(),
                    "page_num": current_page
                })
        elif ext in IMAGE_EXTENSIONS:
            content = extract_text_from_image(file_path)
            sections = [{"section_name": "Full Document", "text_content": content, "page_num": 1}]
        else:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            sections = [{"section_name": "Full Document", "text_content": content, "page_num": 1}]

        # If regex didn't find clear sections (e.g. less than 2 sections parsed)
        # fallback to using RecursiveCharacterTextSplitter chunks as sections
        if len(sections) <= 2:
            json_parsed = parse_document_to_json(file_path)
            sections = []
            for ch in json_parsed["chunks"]:
                sections.append({
                    "section_name": f"Section {ch['metadata']['chunk_index'] + 1}",
                    "text_content": ch["text_content"],
                    "page_num": None
                })
        return sections
    except Exception as e:
        # Full fallback to RecursiveCharacterTextSplitter
        print(f"Regex parser failed: {e}. Falling back to RecursiveCharacterTextSplitter...")
        json_parsed = parse_document_to_json(file_path)
        sections = []
        for ch in json_parsed["chunks"]:
            sections.append({
                "section_name": f"Section {ch['metadata']['chunk_index'] + 1}",
                "text_content": ch["text_content"],
                "page_num": None
            })
        return sections


def enforce_chunk_bounds(sections: List[Dict[str, Any]], max_chars: int = 1000,
                          chunk_overlap: int = 150) -> List[Dict[str, Any]]:
    """Splits any regex-derived section exceeding max_chars into sub-chunks
    via RecursiveCharacterTextSplitter (Stage 1 chunking, no LLM). Sections
    at or under the bound pass through unchanged — short clauses are
    legitimate and shouldn't be padded or merged. This closes the gap where
    parse_document()'s regex sections could otherwise be arbitrarily long."""
    splitter = RecursiveCharacterTextSplitter(chunk_size=max_chars, chunk_overlap=chunk_overlap, length_function=len)
    bounded = []
    for sec in sections:
        text = sec.get("text_content", "")
        if len(text) <= max_chars:
            bounded.append(sec)
            continue

        parts = splitter.split_text(text)
        total = len(parts)
        base_name = sec.get("section_name", "Section")
        for idx, part in enumerate(parts):
            bounded.append({
                **sec,
                "section_name": f"{base_name} (part {idx + 1}/{total})",
                "text_content": part,
                "parent_section_name": base_name,
                "chunk_index": idx,
            })
    return bounded


def parse_document_pages(file_path: str) -> List[Dict[str, Any]]:
    """Extracts raw per-page text (Stage 1, no LLM) for the new `pages`
    Mongo collection. PDF only — python-docx has no reliable page concept,
    so DOCX/TXT sources return [] (an honest reflection of DOCX having no
    fixed pagination, not a bug). Additive: parse_document()'s return shape
    and every existing caller are unaffected."""
    ext = os.path.splitext(file_path)[1].lower()
    if ext != ".pdf":
        return []

    pages = []
    with pdfplumber.open(file_path) as pdf:
        for page_idx, page in enumerate(pdf.pages):
            pages.append({
                "page_number": page_idx + 1,
                "raw_text": page.extract_text() or "",
            })
    return pages
